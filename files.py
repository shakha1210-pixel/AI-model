"""
files.py — Foydalanuvchi yuklagan faylni matn ko'rinishiga o'giradi, shu
matn keyin agentga (Claude/Gemini) suhbat xabari ichida yuboriladi.

Qo'llab-quvvatlanadigan formatlar:
  - Matnli fayllar (.txt, .md, .html, .css, .json, .csv, .yml, kod
    fayllari va h.k.) — to'g'ridan-to'g'ri UTF-8 sifatida o'qiladi.
  - .docx (Word) — python-docx orqali paragraf va jadval matni ajratiladi.
  - .xlsx (Excel) — openpyxl orqali har bir varaq jadval ko'rinishida
    matnga aylantiriladi (juda katta varaqlar qisqartiriladi).

Eski binar formatlar (.doc, .xls) qo'llab-quvvatlanmaydi — ularni to'g'ri
o'qish uchun tashqi dastur (masalan LibreOffice) kerak bo'lardi, bu
loyiha ko'lami uchun ortiqcha murakkablik hisoblanadi. Foydalanuvchiga
buni aniq tushuntiruvchi xabar qaytariladi.
"""

import io

from docx import Document
from fastapi import HTTPException
from openpyxl import load_workbook

MAX_EXTRACTED_CHARS = 8000
MAX_ROWS_PER_SHEET = 200

TEXT_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".css", ".json", ".md",
    ".txt", ".csv", ".yml", ".yaml", ".sql", ".sh", ".java", ".c", ".cpp",
    ".go", ".rs", ".rb", ".php", ".xml",
}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx"}
UNSUPPORTED_HINT = {
    ".doc": "Eskirgan .doc formati qo'llab-quvvatlanmaydi — faylni Word'da "
            "\".docx\" sifatida qayta saqlab, shuni yuklang.",
    ".xls": "Eskirgan .xls formati qo'llab-quvvatlanmaydi — faylni Excel'da "
            "\".xlsx\" sifatida qayta saqlab, shuni yuklang.",
}


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"# Varaq: {sheet.title}")
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i >= MAX_ROWS_PER_SHEET:
                parts.append(f"... ({sheet.title} varag'ida yana qatorlar bor, qisqartirildi)")
                break
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    """Fayl baytlarini o'qiladigan matnga aylantiradi. Qo'llab-quvvatlanmaydigan
    yoki bo'sh natija bo'lsa, aniq sababli HTTPException(400) ko'taradi."""
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in UNSUPPORTED_HINT:
        raise HTTPException(status_code=400, detail=UNSUPPORTED_HINT[ext])

    try:
        if ext in DOCX_EXTENSIONS:
            text = _extract_docx(data)
        elif ext in XLSX_EXTENSIONS:
            text = _extract_xlsx(data)
        elif ext in TEXT_EXTENSIONS:
            text = data.decode("utf-8", errors="replace")
        else:
            raise HTTPException(status_code=400, detail=f"'{ext}' fayl turi qo'llab-quvvatlanmaydi.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Faylni o'qib bo'lmadi: {exc}") from exc

    if not text.strip():
        raise HTTPException(status_code=400, detail="Faylda o'qiladigan matn topilmadi.")

    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS] + "\n... (fayl qisqartirildi)"
    return text
